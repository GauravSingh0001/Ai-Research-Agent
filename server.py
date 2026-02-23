"""
AI Paper Reviewer — Flask API Server
Bridges the dashboard frontend to the Python AI backend.
Last updated: 2026-02-23 - Fixed werkezeug compatibility for Vercel

Endpoints:
  GET    /api/status              — health check + current state
  POST   /api/search              — search papers via Semantic Scholar
  GET    /api/papers              — load saved papers from data/papers.json
  POST   /api/pipeline/run        — run analysis pipeline on saved papers
  GET    /api/pipeline/status     — get pipeline stage statuses
  GET    /api/synthesis           — load latest synthesis (md + sections)
  POST   /api/synthesis/run       — trigger full synthesis generation
  POST   /api/synthesis/revise    — revise synthesis with user instruction
  GET    /api/reports             — list all archived output folders
  GET    /api/reports/<folder>    — get a specific report's data
  GET    /api/export/apa          — download APA references as .txt
  GET    /api/export/bib          — download BibTeX file
  GET    /api/export/pdf          — download PDF of synthesis
  POST   /api/chat                — AI chat about loaded papers/synthesis
  DELETE /api/chat                — clear chat history
"""

import sys
import json
import threading
import uuid
import time
import re
import io
import os
from pathlib import Path
from datetime import datetime
from flask import Flask, jsonify, request, send_file, abort, send_from_directory
from flask_cors import CORS

# ── Path setup ────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(BASE_DIR))

from src.config import (
    PAPERS_DATA_FILE,
    ANALYSIS_RESULTS_FILE,
    RESEARCH_SYNTHESIS_FILE,
    SECTIONS_DATA_FILE,
    BIBTEX_FILE,
    OUTPUT_DIR,
    DATA_DIR,
)
from src.search import search_papers, save_papers, search_multiple_topics, save_cleaned_dataset
from src.analysis import PaperAnalyzer, SIMILARITY_FILE, SECTIONS_DIR
from src.writing import ResearchWriter
from src.utils import setup_logger

logger = setup_logger(__name__)

# Deployment Verification Marker
logger.info("--- STARTING VERCEL DEPLOYMENT (Fix Applied) ---")


app = Flask(__name__, static_folder='dashboard', template_folder='dashboard')
# Read allowed origins from env var for production flexibility
# Default allows all origins for local development and Render deployment
_cors_origins = os.getenv("CORS_ORIGINS", "*").split(",")
CORS(app, origins=_cors_origins, supports_credentials=True)
if "*" in _cors_origins:
    logger.warning("[SERVER] CORS is open to all origins. Set CORS_ORIGINS env var in production.")

# ── In-memory pipeline state ──────────────────────────────
pipeline_state = {
    "running": False,
    "stages": [
        {"id": "pdf-parse",       "title": "PDF Parsing",              "subtitle": "Waiting…", "status": "pending", "progress": 0},
        {"id": "section-extract", "title": "Section Extraction",       "subtitle": "Waiting…", "status": "pending", "progress": 0},
        {"id": "key-findings",    "title": "Key Finding Identification","subtitle": "Waiting…", "status": "pending", "progress": 0},
        {"id": "cross-compare",   "title": "Cross-Paper Comparison",   "subtitle": "Waiting…", "status": "pending", "progress": 0},
        {"id": "embedding",       "title": "Semantic Embedding",       "subtitle": "Waiting…", "status": "pending", "progress": 0},
        {"id": "synthesis-queue", "title": "Synthesis Queue",          "subtitle": "Waiting…", "status": "pending", "progress": 0},
    ],
    "last_run": None,
    "error": None,
}

synthesis_state = {
    "running": False,
    "done": False,
    "error": None,
    "last_run": None,
    "generation_id": None,  # Track which synthesis run is active
    "started_at": None,  # Track when synthesis started
    "revised_markdown": None,  # Store the revised markdown for revision polling
    "output_sections": {},    # Store sections like critique, suggestions
    "generated_files": [],     # Store links to PDFs or other reports
}

# Thread safety lock for all shared state mutations
_state_lock = threading.Lock()

# ── Helpers ───────────────────────────────────────────────
def _reset_pipeline():
    with _state_lock:
        for s in pipeline_state["stages"]:
            s["status"] = "pending"
            s["progress"] = 0
            s["subtitle"] = "Waiting…"
        pipeline_state["running"] = False
        pipeline_state["error"] = None

def _reset_synthesis():
    """Reset synthesis state to initial values."""
    with _state_lock:
        synthesis_state["running"] = False
        synthesis_state["done"] = False
        synthesis_state["error"] = None
        synthesis_state["generation_id"] = None
        synthesis_state["started_at"] = None
        synthesis_state["revised_markdown"] = None
        synthesis_state["output_sections"] = {}
        synthesis_state["generated_files"] = []
        synthesis_state["output_sections"] = {}
        synthesis_state["generated_files"] = []

def _set_stage(idx, status, subtitle, progress):
    with _state_lock:
        s = pipeline_state["stages"][idx]
        s["status"] = status
        s["subtitle"] = subtitle
        s["progress"] = progress

def _load_json(path: Path):
    if not path.exists():
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None

def _latest_output_folder():
    """Return the most recently modified output subfolder, or None if none exist."""
    if not OUTPUT_DIR.exists() or not OUTPUT_DIR.is_dir():
        return None
    try:
        folders = [d for d in OUTPUT_DIR.iterdir() if d.is_dir()]
    except OSError:
        return None
    if not folders:
        return None
    return max(folders, key=lambda d: d.stat().st_mtime)

def _papers_to_apa(papers):
    """Convert papers list to APA 7th edition plain text."""
    lines = []
    for p in papers:
        authors_raw = p.get("authors", [])
        if isinstance(authors_raw, list):
            authors_str = ", ".join(authors_raw)
        else:
            authors_str = str(authors_raw)
        year   = p.get("year", "n.d.")
        title  = p.get("title", "Untitled")
        venue  = p.get("venue", "")
        url    = p.get("url", "")
        line = f"{authors_str} ({year}). {title}."
        if venue:
            line += f" *{venue}*."
        if url:
            line += f" {url}"
        lines.append(line)
    return "\n\n".join(lines)

# ── Routes ────────────────────────────────────────────────

# ── Static File Serving ────────────────────────────────────
DASHBOARD_DIR = BASE_DIR / "dashboard"
if not DASHBOARD_DIR.exists():
    # Fall back to serving files from project root if dashboard/ doesn't exist
    DASHBOARD_DIR = BASE_DIR
    logger.warning(
        f"[SERVER] 'dashboard/' folder not found. Serving static files from project root: {BASE_DIR}\n"
        f"         Move index.html, app.js, styles.css into a 'dashboard/' subfolder to suppress this warning."
    )

@app.route('/')
def serve_index():
    """Serve the main dashboard index.html"""
    return send_from_directory(str(DASHBOARD_DIR), 'index.html')

@app.route('/<path:filename>')
def serve_static(filename):
    """Serve static files (CSS, JS, etc.) from dashboard folder"""
    try:
        return send_from_directory(str(DASHBOARD_DIR), filename)
    except Exception:
        abort(404)

# ── API Routes ────────────────────────────────────────────

@app.route("/api/status")
def api_status():
    papers = _load_json(PAPERS_DATA_FILE)
    synthesis_exists = RESEARCH_SYNTHESIS_FILE.exists()
    return jsonify({
        "ok": True,
        "papers_loaded": len(papers) if papers else 0,
        "synthesis_ready": synthesis_exists,
        "pipeline_running": pipeline_state["running"],
        "synthesis_running": synthesis_state["running"],
        "timestamp": datetime.now().isoformat(),
    })


@app.route("/api/search", methods=["POST"])
def api_search():
    body  = request.get_json(force=True, silent=True) or {}
    topics_raw = body.get("topics") or body.get("topic") or ""
    limit = int(body.get("limit", 5))

    # Accept either a single topic string or a list
    if isinstance(topics_raw, list):
        topics = [t.strip() for t in topics_raw if t.strip()]
    else:
        topics = [t.strip() for t in str(topics_raw).split(",") if t.strip()]

    if not topics:
        return jsonify({"error": "topic is required"}), 400

    try:
        if len(topics) == 1:
            results = search_papers(topics[0], limit=limit)
            dataset = {topics[0]: results}
        else:
            dataset = search_multiple_topics(topics, limit_per_topic=limit)
            results = [p for papers in dataset.values() for p in papers]

        if results:
            save_cleaned_dataset(dataset)
        return jsonify({"papers": results, "count": len(results), "topics": topics, "dataset": dataset})
    except Exception as e:
        logger.error(f"Search error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/papers")
def api_papers():
    papers = _load_json(PAPERS_DATA_FILE)
    if papers is None:
        return jsonify({"papers": [], "count": 0})
    return jsonify({"papers": papers, "count": len(papers)})


@app.route("/api/dataset")
def api_dataset():
    """Return the multi-topic cleaned_dataset.json."""
    cleaned = _load_json(DATA_DIR / "cleaned_dataset.json")
    if cleaned is None:
        return jsonify({"dataset": {}, "topics": [], "total": 0})
    total = sum(len(v) for v in cleaned.values())
    return jsonify({"dataset": cleaned, "topics": list(cleaned.keys()), "total": total})


@app.route("/api/similarity")
def api_similarity():
    """Return the cosine similarity matrix from similarity_results.json."""
    data = _load_json(SIMILARITY_FILE)
    if data is None:
        # Try latest archived folder
        folder = _latest_output_folder()
        if folder:
            data = _load_json(folder / "similarity_results.json")
    if data is None:
        return jsonify({"error": "No similarity results. Run extraction pipeline first."}), 404
    return jsonify(data)


@app.route("/api/sections")
def api_sections():
    """List all section text files available."""
    if not SECTIONS_DIR.exists():
        return jsonify({"sections": [], "papers": []})
    papers_sections = []
    for paper_dir in sorted(SECTIONS_DIR.iterdir()):
        if paper_dir.is_dir():
            files = {f.stem: f.read_text(encoding='utf-8') for f in paper_dir.glob('*.txt')}
            papers_sections.append({"paper": paper_dir.name, "sections": files})
    return jsonify({"papers": papers_sections, "count": len(papers_sections)})


# ── Pipeline ──────────────────────────────────────────────

def _run_pipeline_thread():
    """Background thread: runs analysis then marks stages done."""
    try:
        papers = _load_json(PAPERS_DATA_FILE) or []
        n = len(papers)

        # Stage 0 — PDF Parsing (simulated; abstracts already loaded)
        _set_stage(0, "running", f"Processing {n} documents…", 30)
        time.sleep(0.5)
        _set_stage(0, "done", f"{n} documents processed", 100)

        # Stage 1 — Section Extraction
        _set_stage(1, "running", "Extracting sections…", 20)
        analyzer = PaperAnalyzer()
        time.sleep(0.3)
        _set_stage(1, "done", "Abstract, Methods, Results, References", 100)

        # Stage 2 — Key Findings
        _set_stage(2, "running", "Identifying findings…", 40)
        results = analyzer.run_analysis()
        total_findings = sum(len(p.get("key_findings", [])) for p in results.get("papers", []))
        _set_stage(2, "done", f"{total_findings} findings extracted", 100)

        # Stage 3 — Cross-Paper Comparison
        _set_stage(3, "running", "Computing similarity matrix…", 60)
        analyzer._save_analysis()
        time.sleep(0.3)
        _set_stage(3, "done", "Similarity matrix computed", 100)

        # Stage 4 — Embedding (simulated)
        _set_stage(4, "running", "Updating vector store…", 50)
        time.sleep(0.4)
        _set_stage(4, "done", "Vector store updated", 100)

        # Stage 5 — Synthesis Queue
        _set_stage(5, "running", "Ready for AI generation", 72)
        time.sleep(0.2)
        _set_stage(5, "done", "Ready for AI generation", 100)

        pipeline_state["last_run"] = datetime.now().isoformat()

    except Exception as e:
        logger.error(f"Pipeline error: {e}")
        pipeline_state["error"] = str(e)
        for s in pipeline_state["stages"]:
            if s["status"] == "running":
                s["status"] = "pending"
    finally:
        pipeline_state["running"] = False


@app.route("/api/pipeline/run", methods=["POST"])
def api_pipeline_run():
    if pipeline_state["running"]:
        return jsonify({"error": "Pipeline already running"}), 409
    if not PAPERS_DATA_FILE.exists():
        return jsonify({"error": "No papers found. Run a search first."}), 400

    _reset_pipeline()
    pipeline_state["running"] = True
    t = threading.Thread(target=_run_pipeline_thread, daemon=True)
    t.start()
    return jsonify({"ok": True, "message": "Pipeline started"})


@app.route("/api/pipeline/status")
def api_pipeline_status():
    return jsonify({
        "running": pipeline_state["running"],
        "stages":  pipeline_state["stages"],
        "last_run": pipeline_state["last_run"],
        "error":   pipeline_state["error"],
    })


# ── Synthesis ─────────────────────────────────────────────

def _run_synthesis_thread(gen_id):
    try:
        writer = ResearchWriter()
        # generate_complete_document() already runs critique + final_report bundling internally
        writer.generate_complete_document()

        # Only mark as done if this is still the active generation
        with _state_lock:
            if synthesis_state["generation_id"] == gen_id:
                synthesis_state["done"] = True
                synthesis_state["last_run"] = datetime.now().isoformat()
                # Copy sections to state for the API
                synthesis_state["output_sections"] = writer.output_sections.copy()

                # Add PDF path to generated files
                synthesis_state["generated_files"] = [
                    {"name": "Full Synthesis Report (PDF)", "url": "/api/export/pdf", "type": "pdf"},
                    {"name": "Synthesis Markdown", "url": "/api/export/markdown", "type": "md"},
                    {"name": "BibTeX References", "url": "/api/export/bib", "type": "bib"},
                    {"name": "APA 7th References", "url": "/api/export/apa", "type": "txt"},
                ]
    except Exception as e:
        logger.error(f"Synthesis error: {e}")
        with _state_lock:
            if synthesis_state["generation_id"] == gen_id:
                synthesis_state["error"] = str(e)
    finally:
        with _state_lock:
            if synthesis_state["generation_id"] == gen_id:
                synthesis_state["running"] = False


@app.route("/api/synthesis/run", methods=["POST"])
def api_synthesis_run():
    if synthesis_state["running"]:
        return jsonify({"error": "Synthesis already running"}), 409
    if not ANALYSIS_RESULTS_FILE.exists():
        return jsonify({"error": "Run the extraction pipeline first."}), 400

    # Generate unique ID for this synthesis run to prevent stale data issues
    gen_id = str(uuid.uuid4())
    with _state_lock:
        synthesis_state["generation_id"] = gen_id
        synthesis_state["running"] = True
        synthesis_state["done"]    = False
        synthesis_state["error"]   = None
        synthesis_state["started_at"] = datetime.now().isoformat()
    t = threading.Thread(target=_run_synthesis_thread, args=(gen_id,), daemon=True)
    t.start()
    return jsonify({"ok": True, "message": "Synthesis started", "generation_id": gen_id})


@app.route("/api/synthesis")
def api_synthesis():
    """Return the latest synthesis content + sections."""
    # Try latest archived folder first, then root output
    folder = _latest_output_folder()
    md_path  = (folder / "Research_Synthesis.md") if folder else None
    sec_path = (folder / "Sections_Data.json")    if folder else None

    # Fallback to root output files
    if not md_path or not md_path.exists():
        md_path  = RESEARCH_SYNTHESIS_FILE
    if not sec_path or not sec_path.exists():
        sec_path = SECTIONS_DATA_FILE

    md_text  = md_path.read_text(encoding="utf-8")  if md_path  and md_path.exists()  else ""
    sections = _load_json(sec_path) if sec_path and sec_path.exists() else {}

    # Parse markdown sections for the dashboard
    parsed = _parse_synthesis_md(md_text)

    # Check if synthesis has actually completed by looking at file modification time
    # vs last_run timestamp to prevent showing stale "done" status
    has_content = bool(md_text and len(md_text) > 100)
    
    # Calculate elapsed time for running synthesis
    elapsed_seconds = 0
    if synthesis_state["running"] and synthesis_state["started_at"]:
        try:
            started = datetime.fromisoformat(synthesis_state["started_at"])
            elapsed = datetime.now() - started
            elapsed_seconds = int(elapsed.total_seconds())
        except:
            pass
    
    # synthesis_ready indicates content is ready for display (done and has content)
    synthesis_ready = synthesis_state["done"] and has_content and not synthesis_state["running"]
    
    return jsonify({
        "markdown":  md_text,
        "sections":  sections,
        "parsed":    parsed,
        "running":   synthesis_state["running"],
        "done":      synthesis_state["done"] and has_content,
        "synthesis_ready": synthesis_ready,  # FIX: Add clear readiness flag for polling
        "error":     synthesis_state["error"],
        "last_run":  synthesis_state["last_run"],
        "generation_id": synthesis_state["generation_id"],
        "revised_markdown": synthesis_state["revised_markdown"],
        "output_sections": synthesis_state["output_sections"],
        "generated_files": synthesis_state["generated_files"],
        "elapsed_seconds": elapsed_seconds,
    })


@app.route("/api/synthesis/revise", methods=["POST"])
def api_synthesis_revise():
    body = request.get_json(force=True, silent=True) or {}
    instruction = body.get("instruction", "").strip()
    
    if not instruction:
        return jsonify({"error": "Instruction required"}), 400
        
    def _revise_thread(instr, gen_id):
        try:
            # Initialize ResearchWriter explicitly with analysis data
            writer = ResearchWriter(analysis_file=ANALYSIS_RESULTS_FILE)
            
            # Load previously generated sections to maintain context
            sec_path = SECTIONS_DATA_FILE
            if sec_path.exists():
                sections_data = _load_json(sec_path)
                if sections_data and isinstance(sections_data, dict):
                    writer.output_sections.update(sections_data)
                    logger.info(f"[REVISE] Loaded {len(sections_data)} sections for context preservation")
            
            # Perform revision (saves document internally)
            revised_markdown = writer.revise_document(instr)
            
            # Re-parse revised content for the dashboard
            revised_parsed = _parse_synthesis_md(revised_markdown)
            
            # Thread-safe update of synthesis state
            with _state_lock:
                if synthesis_state["generation_id"] == gen_id:
                    synthesis_state["revised_markdown"] = revised_markdown
                    synthesis_state["output_sections"] = writer.output_sections.copy()
                    synthesis_state["done"] = True
                    synthesis_state["last_run"] = datetime.now().isoformat()
                    # We can also put the parsed result here if needed, 
                    # but api_synthesis parses it on the fly from the file.
                    # HOWEVER, since files are written mid-thread, we ensure safety.
        except Exception as e:
            logger.error(f"Revision error: {e}")
            with _state_lock:
                if synthesis_state["generation_id"] == gen_id:
                    synthesis_state["error"] = str(e)
        finally:
            with _state_lock:
                if synthesis_state["generation_id"] == gen_id:
                    synthesis_state["running"] = False

    if synthesis_state["running"]:
        return jsonify({"error": "Synthesis/Revision already in progress"}), 409

    gen_id = str(uuid.uuid4())
    with _state_lock:
        synthesis_state["generation_id"] = gen_id
        synthesis_state["running"] = True
        synthesis_state["done"] = False
        synthesis_state["error"] = None
        synthesis_state["started_at"] = datetime.now().isoformat()
    t = threading.Thread(target=_revise_thread, args=(instruction, gen_id), daemon=True)
    t.start()
    
    return jsonify({"ok": True, "message": "Revision started", "generation_id": gen_id})


def _parse_synthesis_md(md: str) -> dict:
    """Extract key sections from the synthesis markdown."""
    def extract(heading_pattern):
        m = re.search(heading_pattern + r'\s*\n+([\s\S]*?)(?=\n## |\Z)', md, re.IGNORECASE)
        return m.group(1).strip() if m else ""

    abstract_raw = extract(r'## Abstract')
    # Strip the "Title:" and "Abstract:" labels if present
    abstract_raw = re.sub(r'^Title:.*\n', '', abstract_raw).strip()
    abstract_raw = re.sub(r'^Abstract:\s*\n', '', abstract_raw).strip()

    return {
        "topic":       re.search(r'\*\*Topic:\*\*\s*(.+)', md).group(1).strip() if re.search(r'\*\*Topic:\*\*\s*(.+)', md) else "",
        "date":        re.search(r'\*\*Generated:\*\*\s*(.+)', md).group(1).strip() if re.search(r'\*\*Generated:\*\*\s*(.+)', md) else "",
        "paper_count": re.search(r'\*\*Papers Reviewed:\*\*\s*(\d+)', md).group(1) if re.search(r'\*\*Papers Reviewed:\*\*\s*(\d+)', md) else "0",
        "model":       re.search(r'\*\*AI Provider:\*\*\s*(.+)', md).group(1).strip() if re.search(r'\*\*AI Provider:\*\*\s*(.+)', md) else "",
        "abstract":    abstract_raw,
        "introduction":extract(r'## 1\. Introduction'),
        "methods":     extract(r'## 2\. Methodological Comparison'),
        "results":     extract(r'## 3\. Results Synthesis'),
        "discussion":  extract(r'## 4\. Discussion'),
        "conclusion":  extract(r'## 5\. Conclusion.*'),
        "references":  extract(r'## References'),
    }


# ── Reports ───────────────────────────────────────────────

@app.route("/api/reports")
def api_reports():
    if not OUTPUT_DIR.exists():
        return jsonify({"reports": [], "count": 0})
    folders = sorted(
        [d for d in OUTPUT_DIR.iterdir() if d.is_dir()],
        key=lambda d: d.stat().st_mtime,
        reverse=True
    )
    reports = []
    for folder in folders:
        md_file  = folder / "Research_Synthesis.md"
        bib_file = folder / "References.bib"
        src_file = folder / "Source_Papers.json"

        papers_data = _load_json(src_file) or []
        md_text     = md_file.read_text(encoding="utf-8") if md_file.exists() else ""
        word_count  = len(md_text.split()) if md_text else 0

        # Parse topic and date from markdown header
        topic_m = re.search(r'\*\*Topic:\*\*\s*(.+)', md_text)
        date_m  = re.search(r'\*\*Generated:\*\*\s*(.+)', md_text)
        model_m = re.search(r'\*\*AI Provider:\*\*\s*(.+)', md_text)

        reports.append({
            "id":        folder.name,
            "title":     re.sub(r"_\d{8}_\d{6}$", "", folder.name).replace("_", " "),
            "topic":     topic_m.group(1).strip() if topic_m else folder.name,
            "date":      date_m.group(1).strip()  if date_m  else "Unknown",
            "model":     model_m.group(1).strip()  if model_m else "Unknown",
            "papers":    len(papers_data),
            "words":     word_count,
            "status":    "ready" if md_file.exists() and word_count > 100 else "draft",
            "has_bib":   bib_file.exists(),
        })
    return jsonify({"reports": reports, "count": len(reports)})


@app.route("/api/reports/<folder_name>")
def api_report_detail(folder_name):
    folder = OUTPUT_DIR / folder_name
    if not folder.exists() or not folder.is_dir():
        abort(404)

    md_file  = folder / "Research_Synthesis.md"
    src_file = folder / "Source_Papers.json"
    bib_file = folder / "References.bib"
    sec_file = folder / "Sections_Data.json"

    md_text  = md_file.read_text(encoding="utf-8")  if md_file.exists()  else ""
    papers   = _load_json(src_file) or []
    bib_text = bib_file.read_text(encoding="utf-8") if bib_file.exists() else ""
    sections = _load_json(sec_file) or {}
    parsed   = _parse_synthesis_md(md_text)

    return jsonify({
        "id":       folder_name,
        "markdown": md_text,
        "parsed":   parsed,
        "papers":   papers,
        "bib":      bib_text,
        "sections": sections,
        "apa":      _papers_to_apa(papers),
    })


# ── Export ────────────────────────────────────────────────

@app.route("/api/export/apa")
def api_export_apa():
    """Return APA references as downloadable .txt from the latest run."""
    folder   = _latest_output_folder()
    src_file = (folder / "Source_Papers.json") if folder else PAPERS_DATA_FILE
    papers   = _load_json(src_file) or _load_json(PAPERS_DATA_FILE) or []

    if not papers:
        return jsonify({"error": "No papers found"}), 404

    apa_text = _papers_to_apa(papers)
    apa_bytes = apa_text.encode("utf-8")
    return send_file(
        io.BytesIO(apa_bytes),
        as_attachment=True,
        download_name="references_APA7.txt",
        mimetype="text/plain"
    )


@app.route("/api/export/bib")
def api_export_bib():
    """Return BibTeX file from the latest run."""
    folder   = _latest_output_folder()
    bib_file = (folder / "References.bib") if folder else BIBTEX_FILE
    if not bib_file or not bib_file.exists():
        bib_file = BIBTEX_FILE
    if not bib_file.exists():
        return jsonify({"error": "No BibTeX file found"}), 404
    return send_file(bib_file, as_attachment=True, download_name="references.bib", mimetype="text/plain")


@app.route("/api/export/markdown")
def api_export_markdown():
    """Return the synthesis markdown file."""
    folder  = _latest_output_folder()
    md_file = (folder / "Research_Synthesis.md") if folder else RESEARCH_SYNTHESIS_FILE
    if not md_file or not md_file.exists():
        md_file = RESEARCH_SYNTHESIS_FILE
    if not md_file.exists():
        return jsonify({"error": "No synthesis found"}), 404
    return send_file(md_file, as_attachment=True, download_name="research_synthesis.md", mimetype="text/markdown")


@app.route("/api/export/pdf")
def api_export_pdf():
    """
    Generate and return PDF of the research synthesis with professional Markdown conversion.
    
    Features:
    - Server-side PDF generation using ReportLab
    - Markdown to PDF conversion (headings, bold, italics, lists)
    - Clean academic styling without raw symbols
    - In-memory streaming with proper headers
    """
    # Check if synthesis exists
    folder  = _latest_output_folder()
    md_file = (folder / "Research_Synthesis.md") if folder else RESEARCH_SYNTHESIS_FILE
    if not md_file or not md_file.exists():
        md_file = RESEARCH_SYNTHESIS_FILE
    if not md_file.exists():
        return jsonify({"error": "No synthesis found"}), 404
    
    # Try to generate PDF using reportlab if available
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib import colors
        import re as regex_module
        
        # Read markdown content
        markdown_content = md_file.read_text(encoding='utf-8')
        
        # Create in-memory PDF
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch,
            title="Research Synthesis",
            author="AI Research Agent",
            subject="Academic Research Synthesis"
        )
        
        # Build styles for academic formatting
        styles = getSampleStyleSheet()
        
        # Override default styles
        styles.add(ParagraphStyle(
            name='BodyText',
            fontName='Times-Roman',
            fontSize=11,
            leading=17,  # 1.8x line height
            alignment=4,  # Justified
            spaceAfter=10
        ))
        
        styles.add(ParagraphStyle(
            name='Heading1',
            fontName='Times-Bold',
            fontSize=18,
            leading=22,
            spaceAfter=12,
            textColor=colors.HexColor('#1a1a2e'),
            pageBreakBefore=False,
            spaceBefore=6
        ))
        
        styles.add(ParagraphStyle(
            name='Heading2',
            fontName='Times-Bold',
            fontSize=14,
            leading=16,
            spaceAfter=10,
            textColor=colors.HexColor('#3730a3'),
            pageBreakBefore=False,
            spaceBefore=4
        ))
        
        styles.add(ParagraphStyle(
            name='Heading3',
            fontName='Times-Bold',
            fontSize=12,
            leading=14,
            spaceAfter=8,
            textColor=colors.HexColor('#555555'),
            pageBreakBefore=False
        ))
        
        # Markdown to ReportLab converter
        def convert_markdown_to_reportlab(text):
            """Convert markdown fragments to ReportLab-compatible HTML with tag support."""
            if not text: return ""
            # Escape HTML special chars first (but preserve already-safe context)
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Bold: **text** → <b>text</b>
            text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
            
            # Italic: *text* → <i>text</i>
            text = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', text)
            
            # Italics alt: _text_ → <i>text</i>
            text = re.sub(r'_([^_]+)_', r'<i>\1</i>', text)
            
            # Code: `text` → <font face="Courier">text</font>
            text = re.sub(r'`([^`]+)`', r'<font face="Courier">\1</font>', text)
            
            return text.strip()
        
        # Parse markdown and build story
        story = []
        lines = markdown_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Top-level heading (H1)
            if stripped.startswith('# '):
                title = stripped.replace('# ', '').strip()
                clean_title = convert_markdown_to_reportlab(title)
                para = Paragraph(clean_title, styles['Heading1'])
                para.allowTags = True
                story.append(para)
                story.append(Spacer(1, 0.2*inch))
            
            # Section heading (H2)
            elif stripped.startswith('## '):
                title = stripped.replace('## ', '').strip()
                clean_title = convert_markdown_to_reportlab(title)
                para = Paragraph(clean_title, styles['Heading2'])
                para.allowTags = True
                story.append(para)
                story.append(Spacer(1, 0.15*inch))
            
            # Subsection heading (H3)
            elif stripped.startswith('### '):
                title = stripped.replace('### ', '').strip()
                clean_title = convert_markdown_to_reportlab(title)
                para = Paragraph(clean_title, styles['Heading3'])
                para.allowTags = True
                story.append(para)
                story.append(Spacer(1, 0.1*inch))
            
            # Bullet points
            elif stripped.startswith('- ') or stripped.startswith('* '):
                item = stripped.lstrip('-* ').strip()
                clean_item = convert_markdown_to_reportlab(item)
                para = Paragraph(f"• {clean_item}", styles['BodyText'])
                para.allowTags = True
                story.append(para)
                story.append(Spacer(1, 0.05*inch))
            
            # Numbered list
            elif re.match(r'^\d+\.\s', stripped):
                item = re.sub(r'^\d+\.\s', '', stripped).strip()
                clean_item = convert_markdown_to_reportlab(item)
                para = Paragraph(clean_item, styles['BodyText'])
                para.allowTags = True
                story.append(para)
                story.append(Spacer(1, 0.05*inch))
            
            # Blank line
            elif stripped == '':
                story.append(Spacer(1, 0.08*inch))
            
            # Horizontal rule (---)
            elif stripped == '---' or stripped == '***' or stripped == '___':
                story.append(Spacer(1, 0.1*inch))
            
            # Regular paragraph
            else:
                if stripped:
                    clean_para = convert_markdown_to_reportlab(stripped)
                    # FIX: Use allowTags=True so HTML tags in markdown are rendered properly
                    para = Paragraph(clean_para, styles['BodyText'])
                    para.allowTags = True  # Enable HTML tag rendering
                    story.append(para)
                    story.append(Spacer(1, 0.06*inch))
            
            i += 1
        
        # Build PDF
        doc.build(story)
        
        # Reset buffer position for reading
        pdf_buffer.seek(0)
        
        pdf_size = len(pdf_buffer.getvalue())
        logger.info(f"[PDF] Generated in-memory PDF ({pdf_size} bytes)")
        
        # Track in synthesis state
        with _state_lock:
            pdf_entry = {
                "name": f"Research Synthesis PDF",
                "url": "/api/export/pdf",
                "type": "pdf",
                "size": pdf_size,
                "generated_at": datetime.now().isoformat()
            }
            if pdf_entry not in synthesis_state["generated_files"]:
                synthesis_state["generated_files"].append(pdf_entry)
        
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=f"research_synthesis_{datetime.now().strftime('%Y%m%d')}.pdf",
            mimetype="application/pdf"
        )
    
    except ImportError as e:
        # reportlab not available - return informative error
        logger.warning(f"[PDF] reportlab not available: {e}. Frontend will use browser print fallback.")
        return jsonify({
            "error": "Server-side PDF generation unavailable",
            "message": "Browser print-to-PDF fallback will be used with academic styling",
            "fallback_styling": "Georgia serif, 14px, 1.8 line-height, justified"
        }), 503
    
    except Exception as e:
        # Other errors
        logger.error(f"[PDF] Generation failed: {e}")
        return jsonify({
            "error": "PDF generation error",
            "message": str(e),
            "fallback": "Browser print mechanism will be used"
        }), 503


# ── Quality Assessment ────────────────────────────────────

@app.route("/api/quality")
def api_quality():
    """
    Assess the quality of the current synthesis using Vedanth-inspired
    ContentReviewer metrics: clarity, coherence, academic_tone, completeness,
    citation quality, and overall score.  Returns actionable suggestions.
    """
    # Load synthesis
    folder  = _latest_output_folder()
    md_file = (folder / "Research_Synthesis.md") if folder else RESEARCH_SYNTHESIS_FILE
    if not md_file or not md_file.exists():
        md_file = RESEARCH_SYNTHESIS_FILE
    if not md_file.exists():
        return jsonify({"error": "No synthesis found. Run synthesis first."}), 404

    text = md_file.read_text(encoding="utf-8")

    # ── Basic statistics ──────────────────────────────────
    words     = text.split()
    sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
    word_count    = len(words)
    sentence_count = len(sentences)
    avg_sent_len   = word_count / sentence_count if sentence_count > 0 else 0

    # ── Section presence ─────────────────────────────────
    sections_found = {
        "abstract":    bool(re.search(r'## Abstract', text, re.I)),
        "introduction":bool(re.search(r'## 1?\. Introduction', text, re.I)),
        "methods":     bool(re.search(r'## \d+\. Method', text, re.I)),
        "results":     bool(re.search(r'## \d+\. Results', text, re.I)),
        "discussion":  bool(re.search(r'## \d+\. Discussion', text, re.I)),
        "conclusion":  bool(re.search(r'## \d+\. Conclusion', text, re.I)),
        "references":  bool(re.search(r'## References', text, re.I)),
        "critique":    bool(re.search(r'## Academic Critique|## Critique', text, re.I)),
    }
    section_score = sum(sections_found.values()) / len(sections_found)

    # ── Clarity score (sentence length + transitions) ──
    transition_words = ['however', 'therefore', 'furthermore', 'moreover',
                        'consequently', 'additionally', 'nevertheless']
    trans_count = sum(1 for w in transition_words if w in text.lower())
    clarity = min(1.0, 0.6 + (0.1 if 15 <= avg_sent_len <= 28 else 0) + min(0.3, trans_count * 0.04))

    # ── Academic tone score ──────────────────────────────
    academic_vocab = ['analysis', 'methodology', 'significant', 'findings',
                      'research', 'study', 'empirical', 'synthesis', 'hypothesis',
                      'framework', 'theoretical', 'implications', 'systematic']
    informal_words = ['really', 'very', 'quite', 'pretty', 'sort of', 'kind of', "doesn't", "can't", "won't"]
    acad_count  = sum(1 for w in academic_vocab if w in text.lower())
    inform_count = sum(1 for w in informal_words if w in text.lower())
    academic_tone = max(0.0, min(1.0, 0.65 + min(0.25, acad_count * 0.02) - min(0.15, inform_count * 0.05)))

    # ── Coherence (logical connectors + paragraphs) ──────
    connectors = ['because', 'since', 'therefore', 'thus', 'consequently',
                  'as a result', 'in contrast', 'on the other hand']
    conn_count  = sum(1 for c in connectors if c in text.lower())
    paragraphs  = [p for p in text.split('\n\n') if p.strip()]
    coherence   = min(1.0, 0.65 + min(0.2, conn_count * 0.025) + (0.15 if len(paragraphs) > 5 else 0))

    # ── Citation quality ─────────────────────────────────
    citation_patterns = [r'\(\d{4}\)', r'\[.*?\]', r'\(.*?,\s*\d{4}.*?\)', r'et al\.']
    has_citations = any(re.search(pat, text) for pat in citation_patterns)
    citation_score = 0.9 if has_citations else 0.4

    # ── Completeness ─────────────────────────────────────
    completeness = min(1.0, 0.5 * section_score + 0.3 * min(1.0, word_count / 2000) + 0.2)

    # ── Overall (weighted) ───────────────────────────────
    overall = round(
        clarity * 0.20 + academic_tone * 0.25 + coherence * 0.20 +
        completeness * 0.20 + citation_score * 0.15, 3
    )

    # ── Actionable suggestions ───────────────────────────
    suggestions = []
    if clarity < 0.75:
        suggestions.append({"type": "clarity", "severity": "medium",
            "text": "Add transition words (however, therefore, furthermore) to improve logical flow between sentences."})
    if academic_tone < 0.75:
        suggestions.append({"type": "academic_tone", "severity": "high",
            "text": "Strengthen academic vocabulary: replace informal phrases with precise scientific terminology."})
    if not sections_found["critique"]:
        suggestions.append({"type": "completeness", "severity": "medium",
            "text": "Run Critique & Revise to add an AI-generated academic review section."})
    if not has_citations:
        suggestions.append({"type": "citations", "severity": "high",
            "text": "In-text citations are missing. Ensure all major claims cite authors and years in APA 7th format."})
    if word_count < 1500:
        suggestions.append({"type": "length", "severity": "medium",
            "text": f"Document is short ({word_count} words). Consider expanding the Results Synthesis or Discussion section."})
    if not sections_found["discussion"]:
        suggestions.append({"type": "structure", "severity": "high",
            "text": "Discussion section missing. Add interpretation of results and comparison with prior literature."})
    if coherence < 0.75:
        suggestions.append({"type": "coherence", "severity": "low",
            "text": "Strengthen logical connectors between paragraphs to improve argument flow."})
    if not suggestions:
        suggestions.append({"type": "quality", "severity": "low",
            "text": "Document meets academic quality standards. Consider running a final peer review."})

    return jsonify({
        "word_count":     word_count,
        "sentence_count": sentence_count,
        "avg_sentence_length": round(avg_sent_len, 1),
        "sections_found": sections_found,
        "scores": {
            "clarity":       round(clarity, 3),
            "academic_tone": round(academic_tone, 3),
            "coherence":     round(coherence, 3),
            "completeness":  round(completeness, 3),
            "citations":     round(citation_score, 3),
            "overall":       overall,
        },
        "overall_percent": round(overall * 100, 1),
        "grade": "A" if overall >= 0.85 else "B" if overall >= 0.75 else "C" if overall >= 0.65 else "D",
        "suggestions": suggestions,
    })


@app.route("/api/export/docx")
def api_export_docx():
    """Export synthesis as a Word document (.docx) via python-docx."""
    folder  = _latest_output_folder()
    md_file = (folder / "Research_Synthesis.md") if folder else RESEARCH_SYNTHESIS_FILE
    if not md_file or not md_file.exists():
        md_file = RESEARCH_SYNTHESIS_FILE
    if not md_file.exists():
        return jsonify({"error": "No synthesis found"}), 404

    md_text = md_file.read_text(encoding="utf-8")

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()
        # Page margins
        section = doc.sections[0]
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

        for line in md_text.split('\n'):
            stripped = line.strip()
            if stripped.startswith('# '):
                h = doc.add_heading(stripped[2:], level=0)
                h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=1)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=2)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', stripped):
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style='List Number')
            elif stripped == '---':
                doc.add_paragraph('─' * 60)
            elif stripped:
                p = doc.add_paragraph()
                # Handle bold (**text**) and italic (*text*)
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    else:
                        p.add_run(part)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name=f"research_synthesis_{datetime.now().strftime('%Y%m%d')}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except ImportError:
        # Fallback: styled HTML download
        html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<title>Research Synthesis</title>
<style>body{{font-family:Georgia,serif;max-width:750px;margin:40px auto;line-height:1.8;font-size:14px;color:#1a1a2e}}
h1{{font-size:24px;border-bottom:2px solid #3730a3;padding-bottom:8px}}
h2{{color:#3730a3;font-size:18px}}h3{{color:#555;font-size:15px}}
p{{margin:10px 0;text-align:justify}}</style></head><body>
<pre style="white-space:pre-wrap;font-family:Georgia,serif">{md_text}</pre>
</body></html>"""
        return send_file(
            io.BytesIO(html.encode("utf-8")),
            as_attachment=True,
            download_name="research_synthesis.html",
            mimetype="text/html"
        )
    except Exception as e:
        logger.error(f"[DOCX] Export failed: {e}")
        return jsonify({"error": str(e)}), 500


# ── Chat ──────────────────────────────────────────────────
# In-memory chat history (per server session)
_chat_history = []   # List of {"role": "user"|"assistant", "content": str}
_chat_lock = threading.Lock()

MAX_CHAT_HISTORY = 20  # Keep last N messages for context

def _build_chat_context() -> str:
    """Build a concise research context string from loaded papers + synthesis."""
    context_parts = []

    # Papers context
    try:
        if PAPERS_DATA_FILE.exists():
            papers = json.loads(PAPERS_DATA_FILE.read_text(encoding="utf-8"))
            if papers:
                context_parts.append("## Loaded Research Papers")
                for i, p in enumerate(papers[:6], 1):
                    title = p.get("title", "Unknown")
                    authors = ", ".join(p.get("authors", [])[:3]) if isinstance(p.get("authors"), list) else str(p.get("authors", ""))
                    year = p.get("year", "n.d.")
                    abstract = (p.get("abstract") or "")[:300]
                    findings = "; ".join(p.get("key_findings", [])[:3]) if p.get("key_findings") else ""
                    context_parts.append(
                        f"{i}. **{title}** ({authors}, {year})\n"
                        f"   Abstract: {abstract}…\n"
                        + (f"   Key Findings: {findings}" if findings else "")
                    )
    except Exception:
        pass

    # Synthesis context (first 2000 chars)
    try:
        if RESEARCH_SYNTHESIS_FILE.exists():
            synthesis_text = RESEARCH_SYNTHESIS_FILE.read_text(encoding="utf-8")[:2000]
            context_parts.append("## Current Synthesis Summary")
            context_parts.append(synthesis_text + "…")
    except Exception:
        pass

    return "\n\n".join(context_parts) if context_parts else "No papers or synthesis loaded yet."


@app.route("/api/chat", methods=["POST", "DELETE"])
def api_chat():
    """
    POST: Chat with AI about the loaded papers/synthesis.
          Body: {"message": "your question", "use_context": true}
    DELETE: Clear the chat history.
    """
    global _chat_history

    if request.method == "DELETE":
        with _chat_lock:
            _chat_history = []
        return jsonify({"status": "cleared", "message": "Chat history cleared."})

    # POST — handle new message
    data = request.get_json(force=True, silent=True) or {}
    user_message = (data.get("message") or "").strip()
    use_context = data.get("use_context", True)

    if not user_message:
        return jsonify({"error": "No message provided"}), 400

    # Append user message to history
    with _chat_lock:
        _chat_history.append({"role": "user", "content": user_message})
        # Trim history
        if len(_chat_history) > MAX_CHAT_HISTORY:
            _chat_history = _chat_history[-MAX_CHAT_HISTORY:]
        history_snapshot = list(_chat_history)

    # Build system prompt
    research_context = _build_chat_context() if use_context else ""
    system_prompt = (
        "You are an expert AI research assistant specializing in academic literature analysis. "
        "You help researchers understand, compare, and synthesize academic papers.\n\n"
        "Your responses should be:\n"
        "- Grounded in the paper data provided (do not hallucinate citations)\n"
        "- Concise yet comprehensive\n"
        "- Written in clear academic English\n"
        "- Formatted with markdown when helpful (bullet points, bold key terms)\n\n"
        + (f"## Research Context (Papers & Synthesis)\n{research_context}" if research_context else "")
    )

    # Build conversation prompt from history
    history_text = ""
    for msg in history_snapshot[:-1]:  # exclude the latest user message
        role_label = "User" if msg["role"] == "user" else "Assistant"
        history_text += f"\n{role_label}: {msg['content']}"

    full_prompt = (
        f"{history_text}\n\nUser: {user_message}\n\nAssistant:"
        if history_text
        else f"User: {user_message}\n\nAssistant:"
    )

    # Generate response via AIEngine
    try:
        from src.ai_engine import AIEngine
        ai = AIEngine()
        result = ai.generate(full_prompt, system_prompt=system_prompt, max_tokens=800)
        if result.get("status") == "success" and result.get("text"):
            reply = result["text"].strip()
        else:
            reply = (
                "I'm sorry, I couldn't generate a response right now. "
                f"Error: {result.get('error', 'Unknown AI error')}. "
                "Please ensure your API keys are configured correctly."
            )
    except Exception as e:
        logger.error(f"[CHAT] AI generation failed: {e}")
        reply = f"Chat error: {str(e)}. Please check server logs."

    # Store assistant reply
    with _chat_lock:
        _chat_history.append({"role": "assistant", "content": reply})
        if len(_chat_history) > MAX_CHAT_HISTORY:
            _chat_history = _chat_history[-MAX_CHAT_HISTORY:]

    return jsonify({
        "reply": reply,
        "history_length": len(_chat_history),
        "papers_context_available": PAPERS_DATA_FILE.exists(),
        "synthesis_context_available": RESEARCH_SYNTHESIS_FILE.exists(),
    })


@app.route("/api/chat/history", methods=["GET"])
def api_chat_history():
    """Return the current chat history."""
    with _chat_lock:
        return jsonify({"history": list(_chat_history)})


# ── Run ───────────────────────────────────────────────────
# Gunicorn imports `app` directly (server:app) — app.run() is for local dev only

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    print("\n" + "="*50)
    print("  AI PAPER REVIEWER — API SERVER")
    print(f"  http://localhost:{port}")
    print("="*50 + "\n")
    app.run(host="0.0.0.0", port=port, debug=False, threaded=True)

